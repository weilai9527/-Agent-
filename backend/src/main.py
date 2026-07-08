from __future__ import annotations

import base64
from collections.abc import Iterator
from datetime import datetime, timedelta, timezone
import hashlib
import io
import json
import math
import os
import random
import re
import time
from typing import Any
from urllib.parse import parse_qsl, urlencode, urlparse, urlunparse
from uuid import uuid4

import httpx
from fastapi import Depends, FastAPI, File, HTTPException, Request, Response, UploadFile, WebSocket, WebSocketDisconnect
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import PlainTextResponse, StreamingResponse

from .env import load_env_file, normalize_certificate_env

load_env_file()
normalize_certificate_env()

from .database import all_rows, db, get_database_path, one
from .kimi_followup import KimiFollowupError, analyze_resume, generate_kimi_followup, generate_opening_question
from .qwen_realtime_tts import QwenRealtimeTtsError, get_qwen_realtime_tts_media_type, stream_qwen_realtime_tts
from .qwen_tts import QwenTtsError, get_qwen_tts_media_type, stream_qwen_tts
from .security import (
    create_token,
    hash_password,
    hash_token,
    is_valid_email,
    normalize_email,
    sanitize_user,
    verify_password,
)

app = FastAPI(title="Multi Agent Interview API")

allowed_origins = [
    origin.strip()
    for origin in os.environ.get("FRONTEND_ORIGIN", "http://127.0.0.1:5173").split(",")
    if origin.strip()
]
app.add_middleware(
    CORSMiddleware,
    allow_origins=allowed_origins,
    allow_credentials=True,
    allow_methods=["GET", "POST", "PUT", "PATCH", "DELETE", "OPTIONS"],
    allow_headers=["Content-Type"],
)

SESSION_COOKIE_NAME = "interview_session"
SESSION_MAX_AGE_SECONDS = 60 * 60 * 24 * 7
RESET_MAX_AGE_SECONDS = 60 * 20
LOGIN_WINDOW_SECONDS = 60 * 15
MAX_LOGIN_ATTEMPTS = 8
login_attempts: dict[str, dict[str, int | float]] = {}

PROFILE_TEXT_FIELDS = [
    ("nickname", "昵称", 60, True),
    ("target_role", "目标岗位", 80, False),
    ("experience_level", "经验水平", 40, False),
    ("company_type", "目标公司类型", 80, False),
    ("target_city", "目标城市", 80, False),
    ("expected_salary", "期望薪资", 80, False),
    ("years_of_experience", "工作年限", 40, False),
    ("education_level", "学历背景", 60, False),
    ("skills", "技能标签", 500, False),
    ("project_keywords", "项目关键词", 500, False),
    ("resume_text", "简历文本", 12000, False),
    ("project_experience", "项目经历", 12000, False),
    ("preferred_interview_type", "默认面试类型", 60, False),
    ("preferred_difficulty", "默认难度", 40, False),
    ("preferred_interviewer_style", "面试官风格", 60, False),
]
MAX_RESUME_FILE_BYTES = 10 * 1024 * 1024
INTERVIEW_STATUSES = {"draft", "running", "completed"}
AGENT_STATUSES = {"pending", "active", "completed"}
MESSAGE_SENDER_TYPES = {"agent", "candidate", "system"}
MESSAGE_TYPES = {"question", "answer", "follow_up", "system", "transcript"}
AGENT_QUESTION_MESSAGE_TYPES = {"question", "follow_up"}
INTERVIEW_TEXT_FIELDS = [
    ("target_role", "目标岗位", 80),
    ("experience_level", "经验等级", 40),
    ("interview_type", "面试类型", 60),
    ("company_context", "公司场景", 120),
    ("focus_areas", "重点方向", 500),
    ("difficulty", "难度", 40),
    ("interviewer_style", "面试官风格", 60),
]
AGENT_TEXT_FIELDS = [
    ("agent_name", "Agent 名称", 80),
    ("agent_type", "Agent 类型", 40),
    ("agent_role", "Agent 角色", 500),
    ("strategy", "面试策略", 1000),
]
DEFAULT_AGENT_TEMPLATES = [
    {
        "agent_name": "技术一面 Agent",
        "agent_type": "technical",
        "agent_role": "负责考察候选人的核心技术基础、项目细节和编码思路。",
        "strategy": "先围绕目标岗位确认技术栈，再根据重点方向逐步追问实现细节。",
    },
    {
        "agent_name": "架构二面 Agent",
        "agent_type": "architecture",
        "agent_role": "负责考察系统设计、工程权衡、性能和稳定性意识。",
        "strategy": "从业务场景切入，观察候选人如何拆解模块、设计数据流和处理边界。",
    },
    {
        "agent_name": "HR Agent",
        "agent_type": "hr",
        "agent_role": "负责考察职业动机、沟通表达、团队协作和岗位匹配度。",
        "strategy": "用行为面试问题理解候选人的经历、偏好和稳定性。",
    },
]


def now_iso_after(seconds: int) -> str:
    return (datetime.now(timezone.utc) + timedelta(seconds=seconds)).strftime("%Y-%m-%d %H:%M:%S")


def error(status_code: int, message: str) -> HTTPException:
    return HTTPException(status_code=status_code, detail={"error": message})


@app.exception_handler(HTTPException)
async def http_exception_handler(_request: Request, exc: HTTPException):
    from fastapi.responses import JSONResponse

    if isinstance(exc.detail, dict):
        return JSONResponse(status_code=exc.status_code, content=exc.detail)
    return JSONResponse(status_code=exc.status_code, content={"error": str(exc.detail)})


def json_body(body: dict | None) -> dict:
    return body or {}


def read_text_field(body: dict, field: str, label: str, max_length: int, required: bool = False) -> tuple[Any, str | None]:
    value = str(body.get(field) or "").strip()
    if required and not value:
        return None, f"{label}不能为空。"
    if not value:
        return None, None
    if len(value) > max_length:
        return None, f"{label}不能超过 {max_length} 个字符。"
    return value, None


def read_url_field(body: dict, field: str, label: str, max_length: int, multiline: bool = False) -> tuple[Any, str | None]:
    value = str(body.get(field) or "").strip()
    if not value:
        return None, None
    if len(value) > max_length:
        return None, f"{label}不能超过 {max_length} 个字符。"
    urls = value.split() if multiline else [value]
    for url in urls:
        parsed = urlparse(url)
        if parsed.scheme not in {"http", "https"} or not parsed.netloc:
            return None, f"{label}只支持 http 或 https 链接。"
    return value, None


def parse_profile_input(body: dict, fallback_name: str) -> dict:
    profile: dict[str, Any] = {}
    for field, label, max_length, required in PROFILE_TEXT_FIELDS:
        value, message = read_text_field(body, field, label, max_length, required)
        if message:
            raise error(400, message)
        profile[field] = value

    profile["nickname"] = profile["nickname"] or fallback_name
    profile["avatar_url"], message = read_url_field(body, "avatar_url", "头像链接", 500)
    if message:
        raise error(400, message)
    profile["portfolio_links"], message = read_url_field(body, "portfolio_links", "作品链接", 1000, multiline=True)
    if message:
        raise error(400, message)
    return profile


def parse_interview_input(body: dict, partial: bool = False) -> dict:
    interview: dict[str, Any] = {}
    for field, label, max_length in INTERVIEW_TEXT_FIELDS:
        if partial and field not in body:
            continue
        value, message = read_text_field(body, field, label, max_length, field == "target_role")
        if message:
            raise error(400, message)
        interview[field] = value
    if partial and not interview:
        raise error(400, "请至少提供一个要更新的字段。")
    return interview


def parse_order_index(value: Any) -> int | None:
    if value is None or value == "":
        return None
    if isinstance(value, bool):
        return None
    try:
        number = int(value)
    except (TypeError, ValueError):
        return None
    return number if 0 <= number <= 99 else None


def parse_agent_input(body: dict, partial: bool = False) -> dict:
    agent: dict[str, Any] = {}
    for field, label, max_length in AGENT_TEXT_FIELDS:
        if partial and field not in body:
            continue
        value, message = read_text_field(body, field, label, max_length, field != "strategy")
        if message:
            raise error(400, message)
        agent[field] = value

    if "order_index" in body:
        order_index = parse_order_index(body.get("order_index"))
        if order_index is None:
            raise error(400, "Agent 顺序必须是 0 到 99 之间的整数。")
        agent["order_index"] = order_index

    if "status" in body:
        status = str(body.get("status") or "").strip()
        if status not in AGENT_STATUSES:
            raise error(400, "Agent 状态不正确。")
        agent["status"] = status

    if partial and not agent:
        raise error(400, "请至少提供一个要更新的字段。")
    return agent


def parse_message_input(body: dict) -> dict:
    sender_type = str(body.get("sender_type") or "").strip()
    message_type = str(body.get("message_type") or "").strip()
    content, message = read_text_field(body, "content", "消息内容", 12000, True)
    if sender_type not in MESSAGE_SENDER_TYPES:
        raise error(400, "消息发送方类型不正确。")
    if message_type not in MESSAGE_TYPES:
        raise error(400, "消息类型不正确。")
    if message:
        raise error(400, message)
    transcript_text, message = read_text_field(body, "transcript_text", "语音转写结果", 12000)
    if message:
        raise error(400, message)

    agent_id = str(body.get("agent_id") or "").strip() or None
    if sender_type == "agent" and not agent_id:
        raise error(400, "Agent 消息必须关联 agent_id。")
    if sender_type != "agent" and agent_id:
        raise error(400, "只有 Agent 消息可以关联 agent_id。")
    if sender_type == "candidate" and message_type not in {"answer", "transcript"}:
        raise error(400, "候选人消息只能是 answer 或 transcript。")
    if sender_type == "system" and message_type != "system":
        raise error(400, "系统消息类型必须是 system。")
    if sender_type == "agent" and message_type not in {"question", "follow_up", "system"}:
        raise error(400, "Agent 消息只能是 question、follow_up 或 system。")

    return {
        "agent_id": agent_id,
        "sender_type": sender_type,
        "message_type": message_type,
        "content": content,
        "transcript_text": transcript_text,
    }


def login_attempt_key(request: Request, email: str) -> str:
    host = request.client.host if request.client else "local"
    return f"{host}:{email}"


def is_login_limited(request: Request, email: str) -> bool:
    key = login_attempt_key(request, email)
    record = login_attempts.get(key)
    if not record:
        return False
    if float(record["reset_at"]) <= time.time():
        login_attempts.pop(key, None)
        return False
    return int(record["count"]) >= MAX_LOGIN_ATTEMPTS


def record_failed_login(request: Request, email: str) -> None:
    key = login_attempt_key(request, email)
    now = time.time()
    record = login_attempts.get(key)
    if not record or float(record["reset_at"]) <= now:
        login_attempts[key] = {"count": 1, "reset_at": now + LOGIN_WINDOW_SECONDS}
    else:
        record["count"] = int(record["count"]) + 1


def clear_failed_logins(request: Request, email: str) -> None:
    login_attempts.pop(login_attempt_key(request, email), None)


def create_session(response: Response, user_id: str, user_agent: str | None) -> None:
    token = create_token()
    db.execute(
        """
        INSERT INTO sessions (id, user_id, token_hash, expires_at, user_agent)
        VALUES (?, ?, ?, ?, ?)
        """,
        (str(uuid4()), user_id, hash_token(token), now_iso_after(SESSION_MAX_AGE_SECONDS), user_agent),
    )
    db.commit()
    response.set_cookie(
        SESSION_COOKIE_NAME,
        token,
        httponly=True,
        samesite="lax",
        secure=os.environ.get("COOKIE_SECURE") == "true",
        path="/",
        max_age=SESSION_MAX_AGE_SECONDS,
    )


def clear_session(response: Response) -> None:
    response.delete_cookie(
        SESSION_COOKIE_NAME,
        httponly=True,
        samesite="lax",
        secure=os.environ.get("COOKIE_SECURE") == "true",
        path="/",
    )


def find_user_by_session(request: Request) -> dict | None:
    token = request.cookies.get(SESSION_COOKIE_NAME)
    if not token:
        return None
    return one(
        """
        SELECT users.id, users.email, users.name, users.status, users.created_at, users.last_login_at
        FROM sessions
        JOIN users ON users.id = sessions.user_id
        WHERE sessions.token_hash = ?
          AND sessions.expires_at > CURRENT_TIMESTAMP
          AND users.status = 'normal'
        """,
        (hash_token(token),),
    )


def find_user_by_session_token(token: str | None) -> dict | None:
    if not token:
        return None
    return one(
        """
        SELECT users.id, users.email, users.name, users.status, users.created_at, users.last_login_at
        FROM sessions
        JOIN users ON users.id = sessions.user_id
        WHERE sessions.token_hash = ?
          AND sessions.expires_at > CURRENT_TIMESTAMP
          AND users.status = 'normal'
        """,
        (hash_token(token),),
    )


def require_auth(request: Request) -> dict:
    user = find_user_by_session(request)
    if not user:
        raise error(401, "请先登录。")
    return user


def ensure_profile(user: dict) -> None:
    if one("SELECT id FROM profiles WHERE user_id = ?", (user["id"],)):
        return
    db.execute("INSERT INTO profiles (id, user_id, nickname) VALUES (?, ?, ?)", (str(uuid4()), user["id"], user["name"]))
    db.commit()


def find_profile_by_user_id(user_id: str) -> dict | None:
    return one(
        """
        SELECT id, user_id, nickname, avatar_url, target_role, experience_level, company_type,
               target_city, expected_salary, years_of_experience, education_level, skills,
               project_keywords, resume_text, project_experience, portfolio_links,
               preferred_interview_type, preferred_difficulty, preferred_interviewer_style,
               created_at, updated_at
        FROM profiles
        WHERE user_id = ?
        """,
        (user_id,),
    )


def resume_source_text(profile: dict | None) -> str:
    if not profile:
        return ""
    fields = [
        "target_role",
        "experience_level",
        "years_of_experience",
        "education_level",
        "skills",
        "project_keywords",
        "resume_text",
        "project_experience",
    ]
    return "\n".join(str(profile.get(field) or "").strip() for field in fields if str(profile.get(field) or "").strip())


def resume_source_hash(profile: dict | None) -> str:
    return hashlib.sha256(resume_source_text(profile).encode("utf-8")).hexdigest()


def find_resume_analysis_by_user_id(user_id: str) -> dict | None:
    return one(
        """
        SELECT id, user_id, source_hash, analysis_json, provider, status, error_message,
               created_at, updated_at
        FROM resume_analyses
        WHERE user_id = ?
        """,
        (user_id,),
    )


def upsert_resume_analysis(user_id: str, source_hash: str, analysis: dict, provider: str, error_message: str | None = None) -> dict | None:
    existing = find_resume_analysis_by_user_id(user_id)
    payload = json.dumps(analysis, ensure_ascii=False)
    if existing:
        db.execute(
            """
            UPDATE resume_analyses
            SET source_hash = ?, analysis_json = ?, provider = ?, status = ?, error_message = ?,
                updated_at = CURRENT_TIMESTAMP
            WHERE user_id = ?
            """,
            (source_hash, payload, provider, "analyzed", error_message, user_id),
        )
    else:
        db.execute(
            """
            INSERT INTO resume_analyses (
              id, user_id, source_hash, analysis_json, provider, status, error_message
            )
            VALUES (?, ?, ?, ?, ?, ?, ?)
            """,
            (str(uuid4()), user_id, source_hash, payload, provider, "analyzed", error_message),
        )
    db.commit()
    return find_resume_analysis_by_user_id(user_id)


def serialize_resume_analysis(row: dict | None) -> dict | None:
    if not row:
        return None
    analysis = parse_json_object(row.get("analysis_json"), {})
    return {
        "id": row.get("id"),
        "user_id": row.get("user_id"),
        "source_hash": row.get("source_hash"),
        "analysis": analysis,
        "provider": row.get("provider"),
        "status": row.get("status"),
        "error_message": row.get("error_message"),
        "created_at": row.get("created_at"),
        "updated_at": row.get("updated_at"),
    }


def ensure_resume_analysis_for_user(user: dict, force: bool = False) -> dict:
    ensure_profile(user)
    profile = find_profile_by_user_id(user["id"])
    source = resume_source_text(profile)
    if not source:
        raise error(400, "请先在个人资料中填写简历文本、项目经历或技能标签。")
    current_hash = resume_source_hash(profile)
    existing = find_resume_analysis_by_user_id(user["id"])
    if existing and existing.get("source_hash") == current_hash and not force:
        serialized = serialize_resume_analysis(existing)
        if serialized:
            return serialized

    result = analyze_resume(profile or {})
    row = upsert_resume_analysis(
        user["id"],
        current_hash,
        result["analysis"],
        result.get("provider") or "local",
        result.get("error"),
    )
    serialized = serialize_resume_analysis(row)
    if not serialized:
        raise error(500, "简历分析结果保存失败。")
    serialized["fallback"] = bool(result.get("fallback"))
    return serialized


def select_resume_point(analysis: dict, interview: dict, messages: list[dict]) -> dict:
    asked_text = "\n".join(str(message.get("content") or "") for message in messages if message.get("sender_type") == "agent")
    candidates: list[dict[str, Any]] = []
    target_role = str(interview.get("target_role") or "").lower()

    for project in analysis.get("projects", []) if isinstance(analysis.get("projects"), list) else []:
        if not isinstance(project, dict):
            continue
        name = str(project.get("name") or "").strip()
        if not name:
            continue
        if name in asked_text:
            weight = 2
        else:
            joined = " ".join([name, " ".join(project.get("highlights", []) if isinstance(project.get("highlights"), list) else [])]).lower()
            weight = 7 if target_role and target_role in joined else 5
        candidates.append({"type": "project", "value": name, "project": project, "weight": weight})

    for skill in analysis.get("core_skills", []) if isinstance(analysis.get("core_skills"), list) else []:
        value = str(skill).strip()
        if value:
            candidates.append({"type": "skill", "value": value, "weight": 4 if value in asked_text else 6})

    for risk in analysis.get("risk_points", []) if isinstance(analysis.get("risk_points"), list) else []:
        value = str(risk).strip()
        if value:
            candidates.append({"type": "risk", "value": value, "weight": 2 if value in asked_text else 3})

    if not candidates:
        return {"type": "project", "value": "项目经历", "weight": 1}

    weights = [max(1, int(candidate.get("weight") or 1)) for candidate in candidates]
    selected = random.SystemRandom().choices(candidates, weights=weights, k=1)[0]
    return {key: value for key, value in selected.items() if key != "weight"}


def find_interview_by_user_id(interview_id: str, user_id: str) -> dict | None:
    return one(
        """
        SELECT id, user_id, target_role, experience_level, interview_type, company_context,
               focus_areas, difficulty, interviewer_style, status, created_at, updated_at,
               started_at, completed_at
        FROM interview_sessions
        WHERE id = ? AND user_id = ?
        """,
        (interview_id, user_id),
    )


def normalize_status(status: Any) -> str | None:
    value = str(status or "").strip()
    return value if value in INTERVIEW_STATUSES else None


def list_agents_by_interview_id(interview_id: str) -> list[dict]:
    return all_rows(
        """
        SELECT id, interview_id, agent_name, agent_type, agent_role, strategy, order_index,
               status, created_at, updated_at
        FROM interview_agents
        WHERE interview_id = ?
        ORDER BY order_index ASC, created_at ASC
        """,
        (interview_id,),
    )


def find_agent_by_interview_id(agent_id: str, interview_id: str) -> dict | None:
    return one(
        """
        SELECT id, interview_id, agent_name, agent_type, agent_role, strategy, order_index,
               status, created_at, updated_at
        FROM interview_agents
        WHERE id = ? AND interview_id = ?
        """,
        (agent_id, interview_id),
    )


def create_default_agents(interview_id: str) -> None:
    for index, agent in enumerate(DEFAULT_AGENT_TEMPLATES):
        db.execute(
            """
            INSERT INTO interview_agents (
              id, interview_id, agent_name, agent_type, agent_role, strategy, order_index, status
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                str(uuid4()),
                interview_id,
                agent["agent_name"],
                agent["agent_type"],
                agent["agent_role"],
                agent["strategy"],
                index,
                "pending",
            ),
        )


def ensure_default_agents(interview_id: str) -> None:
    existing = one("SELECT COUNT(*) AS count FROM interview_agents WHERE interview_id = ?", (interview_id,))
    if not existing or existing["count"] == 0:
        create_default_agents(interview_id)
        db.commit()


def next_agent_order_index(interview_id: str) -> int:
    row = one("SELECT MAX(order_index) AS max_order_index FROM interview_agents WHERE interview_id = ?", (interview_id,))
    value = row["max_order_index"] if row else None
    return int(value) + 1 if isinstance(value, int) else 0


def list_messages_by_interview_id(interview_id: str) -> list[dict]:
    return all_rows(
        """
        SELECT messages.id, messages.interview_id, messages.agent_id,
               agents.agent_name, agents.agent_type,
               messages.sender_type, messages.message_type, messages.content,
               messages.transcript_text, messages.order_index, messages.created_at
        FROM interview_messages AS messages
        LEFT JOIN interview_agents AS agents ON agents.id = messages.agent_id
        WHERE messages.interview_id = ?
        ORDER BY messages.order_index ASC, messages.created_at ASC
        """,
        (interview_id,),
    )


def find_message_by_interview_id(message_id: str, interview_id: str) -> dict | None:
    return one(
        """
        SELECT messages.id, messages.interview_id, messages.agent_id,
               agents.agent_name, agents.agent_type,
               messages.sender_type, messages.message_type, messages.content,
               messages.transcript_text, messages.order_index, messages.created_at
        FROM interview_messages AS messages
        LEFT JOIN interview_agents AS agents ON agents.id = messages.agent_id
        WHERE messages.id = ? AND messages.interview_id = ?
        """,
        (message_id, interview_id),
    )


def next_message_order_index(interview_id: str) -> int:
    row = one("SELECT MAX(order_index) AS max_order_index FROM interview_messages WHERE interview_id = ?", (interview_id,))
    value = row["max_order_index"] if row else None
    return int(value) + 1 if isinstance(value, int) else 0


def clamp_score(score: float) -> int:
    return max(1, min(100, round(score)))


def clamp_stat_score(score: float) -> int:
    return max(0, min(100, round(score)))


def average(numbers: list[Any], fallback: float = 0) -> float:
    valid = [float(number) for number in numbers if isinstance(number, (int, float)) and math.isfinite(number)]
    return sum(valid) / len(valid) if valid else fallback


def parse_json_object(value: Any, fallback: dict | None = None) -> dict:
    if fallback is None:
        fallback = {}
    try:
        parsed = json.loads(value or "")
        return parsed if isinstance(parsed, dict) else fallback
    except (TypeError, json.JSONDecodeError):
        return fallback


def generate_mock_evaluation(message: dict) -> dict:
    content = str(message.get("content") or "")
    length_score = 18 if len(content) >= 180 else 12 if len(content) >= 80 else 6
    structure_score = 10 if re.search(r"第一|第二|首先|其次|最后|因为|所以|例如|比如", content) else 4
    technical_score = 14 if re.search(r"架构|性能|数据库|缓存|并发|接口|边界|测试|部署|Agent|SQL|Node|React", content, re.I) else 6
    base_score = clamp_score(58 + length_score + structure_score + technical_score)
    expression_clarity = clamp_score(base_score + (4 if structure_score >= 10 else -4))
    technical_depth = clamp_score(base_score + (5 if technical_score >= 14 else -6))
    business_understanding = clamp_score(base_score + (5 if re.search(r"业务|用户|场景|指标|成本|收益|风险", content) else -5))
    return {
        "score": base_score,
        "strengths": "\n".join(
            [
                "回答有一定展开，能够覆盖背景和做法。" if len(content) >= 80 else "回答能抓住问题方向，具备继续追问的基础。",
                "能提到具体技术点，便于面试官继续深挖。" if technical_score >= 14 else "表达保持聚焦，没有明显跑题。",
            ]
        ),
        "issues": "\n".join(
            [
                "可以进一步压缩表达，让重点更突出。" if len(content) >= 180 else "细节还不够充分，关键方案、权衡和结果需要补充。",
                "结构已经初步清晰，但结论和量化结果还可以更靠前。" if structure_score >= 10 else "回答结构略散，建议按背景、行动、结果组织。",
            ]
        ),
        "suggestions": "\n".join(
            [
                "补充一个具体场景，说明你负责的模块、遇到的限制和最终结果。",
                "用 1-2 个量化指标呈现影响，例如耗时、成功率、性能或成本变化。",
                "主动说明方案取舍，展示你对边界条件和风险的判断。",
            ]
        ),
        "dimension_scores": {
            "technical_depth": technical_depth,
            "expression_clarity": expression_clarity,
            "business_understanding": business_understanding,
        },
    }


def find_evaluation_by_message_id(interview_id: str, message_id: str) -> dict | None:
    return one(
        """
        SELECT evaluations.id, evaluations.interview_id, evaluations.message_id, evaluations.agent_id,
               agents.agent_name, agents.agent_type,
               evaluations.score, evaluations.strengths, evaluations.issues, evaluations.suggestions,
               evaluations.dimension_scores, evaluations.created_at, evaluations.updated_at
        FROM interview_evaluations AS evaluations
        LEFT JOIN interview_agents AS agents ON agents.id = evaluations.agent_id
        WHERE evaluations.interview_id = ? AND evaluations.message_id = ?
        """,
        (interview_id, message_id),
    )


def find_evaluation_by_id(interview_id: str, evaluation_id: str) -> dict | None:
    return one(
        """
        SELECT evaluations.id, evaluations.interview_id, evaluations.message_id, evaluations.agent_id,
               agents.agent_name, agents.agent_type,
               evaluations.score, evaluations.strengths, evaluations.issues, evaluations.suggestions,
               evaluations.dimension_scores, evaluations.created_at, evaluations.updated_at
        FROM interview_evaluations AS evaluations
        LEFT JOIN interview_agents AS agents ON agents.id = evaluations.agent_id
        WHERE evaluations.interview_id = ? AND evaluations.id = ?
        """,
        (interview_id, evaluation_id),
    )


def list_evaluations_by_interview_id(interview_id: str) -> list[dict]:
    return all_rows(
        """
        SELECT evaluations.id, evaluations.interview_id, evaluations.message_id, evaluations.agent_id,
               agents.agent_name, agents.agent_type,
               evaluations.score, evaluations.strengths, evaluations.issues, evaluations.suggestions,
               evaluations.dimension_scores, evaluations.created_at, evaluations.updated_at,
               messages.content AS message_content, messages.order_index AS message_order_index
        FROM interview_evaluations AS evaluations
        LEFT JOIN interview_agents AS agents ON agents.id = evaluations.agent_id
        JOIN interview_messages AS messages ON messages.id = evaluations.message_id
        WHERE evaluations.interview_id = ?
        ORDER BY messages.order_index ASC, evaluations.created_at ASC
        """,
        (interview_id,),
    )


def grade_from_score(score: int) -> str:
    if score >= 90:
        return "A"
    if score >= 80:
        return "B"
    if score >= 70:
        return "C"
    if score >= 60:
        return "D"
    return "E"


def pass_recommendation_from_score(score: int) -> str:
    if score >= 85:
        return "strong_pass"
    if score >= 75:
        return "pass"
    if score >= 60:
        return "borderline"
    return "no_pass"


def build_ability_radar(evaluations: list[dict]) -> dict:
    radar = {}
    for dimension in ["technical_depth", "expression_clarity", "business_understanding"]:
        radar[dimension] = clamp_score(average([parse_json_object(e.get("dimension_scores")).get(dimension) for e in evaluations], 70))
    return radar


def build_agent_feedback(agents: list[dict], evaluations: list[dict]) -> list[dict]:
    feedback = []
    for agent in agents:
        related = [evaluation for evaluation in evaluations if evaluation.get("agent_id") == agent["id"]]
        score = clamp_score(average([evaluation["score"] for evaluation in related], 72))
        feedback.append(
            {
                "agent_id": agent["id"],
                "agent_name": agent["agent_name"],
                "agent_type": agent["agent_type"],
                "score": score,
                "comment": (
                    f"已完成 {len(related)} 条回答复盘，整体表现{'稳定' if score >= 80 else '仍需加强'}。"
                    if related
                    else "暂无关联单轮评价，后续可结合该 Agent 的追问补充更细的判断。"
                ),
            }
        )
    return feedback


def build_timeline_review(messages: list[dict], evaluations: list[dict]) -> list[dict]:
    evaluation_by_message_id = {evaluation["message_id"]: evaluation for evaluation in evaluations}
    timeline = []
    for message in messages:
        content = message.get("content") or ""
        timeline.append(
            {
                "message_id": message["id"],
                "order_index": message["order_index"],
                "sender_type": message["sender_type"],
                "message_type": message["message_type"],
                "agent_name": message.get("agent_name"),
                "content_preview": f"{content[:80]}..." if len(content) > 80 else content,
                "score": evaluation_by_message_id.get(message["id"], {}).get("score"),
            }
        )
    return timeline


def generate_mock_report(interview: dict, agents: list[dict], messages: list[dict], evaluations: list[dict]) -> dict:
    total_score = clamp_score(average([evaluation["score"] for evaluation in evaluations], 72 if messages else 60))
    candidate_answers = len([message for message in messages if message["sender_type"] == "candidate"])
    agent_questions = len([message for message in messages if message["sender_type"] == "agent"])
    return {
        "total_score": total_score,
        "grade": grade_from_score(total_score),
        "pass_recommendation": pass_recommendation_from_score(total_score),
        "ability_radar": build_ability_radar(evaluations),
        "agent_feedback": build_agent_feedback(agents, evaluations),
        "timeline_review": build_timeline_review(messages, evaluations),
        "summary": f"本次模拟面向{interview['target_role']}，共记录 {len(messages)} 条消息，其中候选人回答 {candidate_answers} 条，Agent 提问或追问 {agent_questions} 条。综合当前单轮评价，整体等级为 {grade_from_score(total_score)}。",
        "suggestions": "\n".join(
            [
                "继续补充回答中的项目背景、关键决策和量化结果。",
                "针对低分维度安排专项训练，优先复盘被追问但回答不充分的问题。",
                "下一次模拟可以提高难度或增加架构追问，检验方案权衡能力。",
            ]
        ),
    }


def create_initial_question(interview: dict, agent: dict | None) -> str:
    role = interview.get("target_role") or "目标岗位"
    focus = interview.get("focus_areas") or "项目经历"
    agent_name = agent.get("agent_name") if agent else "面试官"
    return f"{agent_name}：请结合你的{focus}，介绍一个最能体现你胜任{role}的项目。"


def build_realtime_session_config(interview: dict, agents: list[dict], messages: list[dict]) -> dict:
    current_agent = next((agent for agent in agents if agent.get("status") == "active"), agents[0] if agents else None)
    latest_questions = [
        message
        for message in messages
        if message["sender_type"] == "agent" and message.get("message_type") in AGENT_QUESTION_MESSAGE_TYPES
    ]
    latest_question = latest_questions[-1]["content"] if latest_questions else create_initial_question(interview, current_agent)
    recent_messages = messages[-8:]

    def speaker_label(message: dict) -> str:
        if message["sender_type"] == "candidate":
            return "候选人"
        if message["sender_type"] == "system":
            return "系统"
        return message.get("agent_name") or "面试官"

    recent_context = "\n".join(
        [
            f"{speaker_label(message)}：{message.get('content') or ''}"
            for message in recent_messages
        ]
    ) or "暂无历史对话。"
    asked_question_items = [message["content"] for message in latest_questions[-5:]]
    asked_questions = "\n".join([f"- {question}" for question in asked_question_items]) or "- 暂无"
    recent_message_items = [
        {
            "sender_type": message["sender_type"],
            "speaker": speaker_label(message),
            "content": message.get("content") or "",
        }
        for message in recent_messages
    ]
    return {
        "type": "realtime",
        "model": os.environ.get("OPENAI_REALTIME_MODEL", "gpt-realtime-2"),
        "instructions": "\n".join(
            [
                "你是一名中文 AI 电话面试官，正在进行一场真实语音模拟面试。",
                "请使用自然、简洁、专业的中文口语交流，每次只问一个问题。",
                "候选人回答后，先做一句短反馈，再围绕项目细节、方案取舍、量化结果或风险边界继续追问。",
                "不要重复已经问过的问题；如果候选人回答泛泛而谈，要换一个角度追问具体案例、数据口径、失败复盘或协作冲突。",
                "连续追问要按层次推进：背景职责 -> 方案取舍 -> 指标验证 -> 线上风险 -> 团队落地。",
                "不要一次性给出长篇评价；把节奏控制成电话面试。",
                f"目标岗位：{interview.get('target_role') or '未填写'}",
                f"面试类型：{interview.get('interview_type') or '综合模拟'}",
                f"难度：{interview.get('difficulty') or '标准'}",
                f"面试官风格：{interview.get('interviewer_style') or '专业追问'}",
                f"当前面试官：{current_agent.get('agent_name') if current_agent else '技术面试 Agent'}",
                f"当前问题：{latest_question}",
                f"最近对话：\n{recent_context}",
                f"已问过的问题，避免复述：\n{asked_questions}",
                "如果刚接通，请从当前问题开始，不要重复介绍系统功能。",
            ]
        ),
        "active_agent": current_agent.get("agent_name") if current_agent else "技术面试 Agent",
        "current_question": latest_question,
        "recent_messages": recent_message_items,
        "asked_questions": asked_question_items,
        "audio": {"output": {"voice": os.environ.get("OPENAI_REALTIME_VOICE", "marin")}},
    }


def build_openai_realtime_session_config(interview: dict, agents: list[dict], messages: list[dict]) -> dict:
    session_config = build_realtime_session_config(interview, agents, messages)
    return {
        key: session_config[key]
        for key in ("type", "model", "instructions", "audio")
        if key in session_config
    }


def log_qwen_tts_metric(event: str, **payload: Any) -> None:
    print(
        "Qwen TTS:",
        json.dumps({"event": event, **payload}, ensure_ascii=False),
        flush=True,
    )


def stream_with_qwen_tts_metrics(
    audio_stream: Iterator[bytes],
    *,
    request_id: str,
    provider: str,
    text_length: int,
    fallback_from: str | None = None,
) -> Iterator[bytes]:
    started_at = time.perf_counter()
    first_packet_ms: int | None = None
    total_bytes = 0
    chunk_count = 0
    status = "completed"
    error_message: str | None = None

    log_qwen_tts_metric(
        "start",
        request_id=request_id,
        provider=provider,
        fallback_from=fallback_from,
        text_length=text_length,
    )

    try:
        for chunk in audio_stream:
            if not chunk:
                continue
            chunk_count += 1
            total_bytes += len(chunk)
            if first_packet_ms is None:
                first_packet_ms = round((time.perf_counter() - started_at) * 1000)
                log_qwen_tts_metric(
                    "first_packet",
                    request_id=request_id,
                    provider=provider,
                    first_packet_ms=first_packet_ms,
                    first_packet_bytes=len(chunk),
                )
            yield chunk
    except Exception as exc:
        status = "failed"
        error_message = str(exc)
        raise
    finally:
        log_qwen_tts_metric(
            "finish",
            request_id=request_id,
            provider=provider,
            fallback_from=fallback_from,
            status=status,
            first_packet_ms=first_packet_ms,
            duration_ms=round((time.perf_counter() - started_at) * 1000),
            chunk_count=chunk_count,
            total_bytes=total_bytes,
            text_length=text_length,
            error=error_message,
        )


def build_qwen_tts_response(text: str, provider_preference: str = "auto") -> StreamingResponse:
    normalized_preference = (provider_preference or "auto").strip().lower()
    if normalized_preference not in {"auto", "realtime", "standard"}:
        raise error(400, "语音合成 provider 只支持 auto、realtime 或 standard。")

    request_id = str(uuid4())
    fallback_from: str | None = None
    fallback_error: str | None = None

    if normalized_preference in {"auto", "realtime"}:
        try:
            audio_stream = stream_qwen_realtime_tts(text)
            provider = "realtime"
            media_type = get_qwen_realtime_tts_media_type()
        except QwenRealtimeTtsError as exc:
            fallback_from = "realtime"
            fallback_error = str(exc)
            log_qwen_tts_metric(
                "fallback",
                request_id=request_id,
                from_provider="realtime",
                to_provider="standard",
                reason=fallback_error,
            )
        else:
            return StreamingResponse(
                stream_with_qwen_tts_metrics(
                    audio_stream,
                    request_id=request_id,
                    provider=provider,
                    text_length=len(text),
                ),
                media_type=media_type,
                headers={"X-Qwen-TTS-Provider": provider, "X-Qwen-TTS-Request-Id": request_id},
            )

    try:
        audio_stream = stream_qwen_tts(text)
    except QwenTtsError as exc:
        message = str(exc)
        if fallback_error:
            message = f"实时合成失败：{fallback_error}；普通合成失败：{message}"
        raise error(500, message)

    provider = "standard"
    headers = {"X-Qwen-TTS-Provider": provider, "X-Qwen-TTS-Request-Id": request_id}
    if fallback_from:
        headers["X-Qwen-TTS-Fallback-From"] = fallback_from
    return StreamingResponse(
        stream_with_qwen_tts_metrics(
            audio_stream,
            request_id=request_id,
            provider=provider,
            fallback_from=fallback_from,
            text_length=len(text),
        ),
        media_type=get_qwen_tts_media_type(),
        headers=headers,
    )


def _qwen_omni_audio_format(mime_type: str) -> str:
    normalized = (mime_type or "").lower()
    if "wav" in normalized:
        return "wav"
    if "mpeg" in normalized or "mp3" in normalized:
        return "mp3"
    if "mp4" in normalized or "m4a" in normalized:
        return "mp4"
    if "ogg" in normalized or "opus" in normalized:
        return "ogg"
    if "webm" in normalized:
        return "webm"
    return os.environ.get("QWEN_OMNI_INPUT_AUDIO_FORMAT", "webm").strip() or "webm"


def _build_qwen_omni_prompt(session_config: dict, start_payload: dict) -> str:
    current_question = str(start_payload.get("currentQuestion") or session_config.get("current_question") or "").strip()
    active_agent = str(start_payload.get("activeAgent") or session_config.get("active_agent") or "").strip()
    asked_questions = session_config.get("asked_questions") or []
    asked_questions_text = "\n".join([f"- {question}" for question in asked_questions[-5:]]) or "- 暂无"
    return "\n".join(
        [
            "候选人刚刚用语音回答了当前问题。请先准确理解这段语音，再继续面试。",
            "绝对不要复述当前问题，不要重新开场，不要介绍系统功能。",
            "如果听清了：先给一句很短的自然反馈，再提出一个新的追问。",
            "如果没有听清：只用一句话请候选人重复或补充。",
            "输出要像真实电话面试：简洁、口语化、每次只问一个问题。",
            f"当前面试官：{active_agent or '技术面试 Agent'}",
            f"当前问题：{current_question or '请围绕候选人的项目经历继续追问。'}",
            f"已问过的问题，避免重复：\n{asked_questions_text}",
        ]
    )


def _build_qwen_omni_messages(session_config: dict, start_payload: dict, audio_data_url: str, input_format: str) -> list[dict]:
    qwen_messages: list[dict] = [
        {
            "role": "system",
            "content": "\n".join(
                [
                    str(session_config.get("instructions") or ""),
                    "现在进入 Qwen-Omni 语音轮次：请根据最后一条用户音频推进面试，不要复述旧问题。",
                ]
            ),
        }
    ]

    recent_messages = session_config.get("recent_messages") or []
    for message in recent_messages[-6:]:
        content = str(message.get("content") or "").strip()
        if not content:
            continue
        qwen_messages.append(
            {
                "role": "user" if message.get("sender_type") == "candidate" else "assistant",
                "content": content[:1200],
            }
        )

    qwen_messages.append(
        {
            "role": "user",
            "content": [
                {"type": "input_audio", "input_audio": {"data": audio_data_url, "format": input_format}},
                {"type": "text", "text": _build_qwen_omni_prompt(session_config, start_payload)},
            ],
        }
    )
    return qwen_messages


def _env_float_value(name: str, default: float) -> float:
    value = os.environ.get(name, "").strip()
    if not value:
        return default
    try:
        return float(value)
    except ValueError as exc:
        raise QwenRealtimeTtsError(f"{name} 必须是数字。") from exc


def _env_int_value(name: str, default: int) -> int:
    value = os.environ.get(name, "").strip()
    if not value:
        return default
    try:
        return int(value)
    except ValueError as exc:
        raise QwenRealtimeTtsError(f"{name} 必须是整数。") from exc


def _env_bool_value(name: str, default: bool = False) -> bool:
    value = os.environ.get(name, "").strip().lower()
    if not value:
        return default
    return value in {"1", "true", "yes", "on"}


def _append_query_param(url: str, key: str, value: str) -> str:
    parsed = urlparse(url)
    query = dict(parse_qsl(parsed.query, keep_blank_values=True))
    query.setdefault(key, value)
    return urlunparse(parsed._replace(query=urlencode(query)))


def _qwen_realtime_model() -> str:
    return os.environ.get("QWEN_OMNI_REALTIME_MODEL", "qwen3.5-omni-plus-realtime")


def _qwen_realtime_api_key() -> str:
    return (
        os.environ.get("QWEN_OMNI_REALTIME_API_KEY")
        or os.environ.get("QWEN_OMNI_API_KEY")
        or os.environ.get("DASHSCOPE_API_KEY")
        or ""
    ).strip()


def _qwen_realtime_webrtc_url(model: str) -> str:
    configured_url = os.environ.get("QWEN_OMNI_REALTIME_WEBRTC_URL", "").strip()
    if configured_url:
        return _append_query_param(configured_url.replace("{model}", model), "model", model)

    configured_endpoint = os.environ.get("QWEN_OMNI_REALTIME_WEBRTC_ENDPOINT", "").strip()
    if not configured_endpoint:
        raise QwenRealtimeTtsError(
            "Qwen-Omni WebRTC 需要配置白名单 Endpoint：请设置 QWEN_OMNI_REALTIME_WEBRTC_ENDPOINT "
            "或完整的 QWEN_OMNI_REALTIME_WEBRTC_URL。普通 dashscope.aliyuncs.com 域名仅用于 WebSocket，"
            "不能直接作为 WebRTC SDP 交换地址。"
        )

    endpoint = configured_endpoint
    if not endpoint.startswith(("http://", "https://")):
        endpoint = f"https://{endpoint}"
    parsed = urlparse(endpoint)
    path = parsed.path.rstrip("/") or "/api/v1/webrtc/realtime"
    if path == "/":
        path = "/api/v1/webrtc/realtime"
    endpoint = urlunparse(parsed._replace(path=path))
    return _append_query_param(endpoint, "model", model)


def build_qwen_realtime_session_update(session_config: dict, start_payload: dict | None = None) -> dict:
    payload = start_payload or {}
    active_agent = str(payload.get("activeAgent") or session_config.get("active_agent") or "").strip()
    current_question = str(payload.get("currentQuestion") or session_config.get("current_question") or "").strip()
    instructions = "\n".join(
        [
            str(session_config.get("instructions") or ""),
            "当前使用 Qwen-Omni-Realtime WebRTC 直连链路。",
            "只有当前面试官 Agent 对候选人说话，后台分析 Agent 不要出声或自我介绍。",
            "如果候选人已经回答，请基于回答继续追问；如果候选人刚接入，请直接提出当前问题。",
            f"当前发声 Agent：{active_agent or session_config.get('active_agent') or '技术面试 Agent'}",
            f"当前问题：{current_question or session_config.get('current_question') or '请围绕候选人的项目经历继续追问。'}",
        ]
    )
    model = _qwen_realtime_model()
    turn_detection_type = os.environ.get("QWEN_OMNI_REALTIME_TURN_DETECTION", "").strip()
    if not turn_detection_type:
        turn_detection_type = "semantic_vad" if "qwen3.5" in model else "server_vad"

    session: dict[str, Any] = {
        "modalities": ["text", "audio"],
        "voice": os.environ.get("QWEN_OMNI_REALTIME_VOICE") or os.environ.get("QWEN_OMNI_VOICE", "Tina"),
        "input_audio_format": "pcm",
        "output_audio_format": "pcm",
        "input_audio_transcription": {
            "model": os.environ.get("QWEN_OMNI_REALTIME_TRANSCRIPTION_MODEL", "qwen3-asr-flash-realtime"),
        },
        "instructions": instructions,
        "smooth_output": _env_bool_value("QWEN_OMNI_REALTIME_SMOOTH_OUTPUT", False),
        "turn_detection": {
            "type": turn_detection_type,
            "threshold": _env_float_value("QWEN_OMNI_REALTIME_VAD_THRESHOLD", 0.5),
            "prefix_padding_ms": _env_int_value("QWEN_OMNI_REALTIME_PREFIX_PADDING_MS", 500),
            "silence_duration_ms": _env_int_value("QWEN_OMNI_REALTIME_SILENCE_DURATION_MS", 800),
        },
    }

    if _env_bool_value("QWEN_OMNI_REALTIME_ENABLE_SEARCH", False):
        session["enable_search"] = True
        session["search_options"] = {"enable_source": _env_bool_value("QWEN_OMNI_REALTIME_SEARCH_SOURCE", True)}

    return {
        "event_id": f"event_{uuid4().hex}",
        "type": "session.update",
        "session": session,
    }


async def _stream_qwen_omni_http_response(websocket: WebSocket, audio_bytes: bytes, mime_type: str, session_config: dict, start_payload: dict) -> None:
    api_key = os.environ.get("QWEN_OMNI_API_KEY") or os.environ.get("DASHSCOPE_API_KEY")
    if not api_key or api_key in {"your-dashscope-api-key", "your-api-key-here"}:
        await websocket.send_json({"type": "error", "message": "请先配置 DASHSCOPE_API_KEY 或 QWEN_OMNI_API_KEY。"})
        return
    if not audio_bytes:
        await websocket.send_json({"type": "error", "message": "没有收到可提交给 Qwen-Omni 的录音。"})
        return

    base_url = os.environ.get("QWEN_OMNI_BASE_URL", "https://dashscope.aliyuncs.com/compatible-mode/v1").rstrip("/")
    model = os.environ.get("QWEN_OMNI_MODEL", "qwen3.5-omni-plus")
    voice = os.environ.get("QWEN_OMNI_VOICE", "Tina")
    audio_format = os.environ.get("QWEN_OMNI_OUTPUT_AUDIO_FORMAT", "wav")
    input_format = os.environ.get("QWEN_OMNI_INPUT_AUDIO_FORMAT", "").strip() or _qwen_omni_audio_format(mime_type)
    audio_base64 = base64.b64encode(audio_bytes).decode("ascii")
    audio_data_url = f"data:audio/{input_format};base64,{audio_base64}"

    stream_response = os.environ.get("QWEN_OMNI_STREAM", "false").lower() == "true"
    request_body = {
        "model": model,
        "messages": _build_qwen_omni_messages(session_config, start_payload, audio_data_url, input_format),
        "modalities": ["text", "audio"],
        "audio": {"voice": voice, "format": audio_format},
        "temperature": float(os.environ.get("QWEN_OMNI_TEMPERATURE", "0.7")),
        "stream": stream_response,
    }
    if stream_response:
        request_body["stream_options"] = {"include_usage": True}

    await websocket.send_json({"type": "status", "message": f"正在调用 Qwen-Omni API：{model}"})
    text_parts: list[str] = []
    audio_chunk_count = 0
    audio_byte_count = 0

    async def send_audio_data(audio_data: str | None) -> None:
        nonlocal audio_chunk_count, audio_byte_count
        if not audio_data:
            return
        clean_data = str(audio_data)
        if "," in clean_data and clean_data.lstrip().startswith("data:"):
            clean_data = clean_data.split(",", 1)[1]
        try:
            audio_byte_count += len(base64.b64decode(clean_data, validate=False))
        except Exception:
            pass
        audio_chunk_count += 1
        await websocket.send_json({"type": "audio_delta", "data": clean_data})

    try:
        async with httpx.AsyncClient(timeout=float(os.environ.get("QWEN_OMNI_TIMEOUT", "120"))) as client:
            if not stream_response:
                response = await client.post(
                    f"{base_url}/chat/completions",
                    headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
                    json=request_body,
                )
                if response.status_code >= 400:
                    message = response.text or "Qwen-Omni API 调用失败。"
                    await websocket.send_json({"type": "error", "message": message[:1200]})
                    return
                data = response.json()
                choices = data.get("choices") or []
                message = (choices[0].get("message") if choices else {}) or {}
                content = message.get("content")
                if content:
                    text_parts.append(str(content))
                    await websocket.send_json({"type": "text", "role": "agent", "text": str(content)})
                audio = message.get("audio") or {}
                transcript = audio.get("transcript")
                if transcript and not content:
                    text_parts.append(str(transcript))
                    await websocket.send_json({"type": "text", "role": "agent", "text": str(transcript)})
                await send_audio_data(audio.get("data"))
            else:
                async with client.stream(
                    "POST",
                    f"{base_url}/chat/completions",
                    headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
                    json=request_body,
                ) as response:
                    if response.status_code >= 400:
                        body = await response.aread()
                        message = body.decode("utf-8", errors="ignore") or "Qwen-Omni API 调用失败。"
                        await websocket.send_json({"type": "error", "message": message[:1200]})
                        return

                    async for line in response.aiter_lines():
                        line = line.strip()
                        if not line or not line.startswith("data:"):
                            continue
                        data = line.removeprefix("data:").strip()
                        if data == "[DONE]":
                            break
                        try:
                            chunk = json.loads(data)
                        except ValueError:
                            continue
                        choices = chunk.get("choices") or []
                        if not choices:
                            continue
                        delta = choices[0].get("delta") or {}
                        content = delta.get("content")
                        if content:
                            text_parts.append(str(content))
                            await websocket.send_json({"type": "text", "role": "agent", "text": str(content)})
                        audio = delta.get("audio") or {}
                        await send_audio_data(audio.get("data"))
                        transcript = audio.get("transcript")
                        if transcript:
                            await websocket.send_json({"type": "text", "role": "agent", "text": str(transcript)})
        if text_parts:
            await websocket.send_json({"type": "transcript", "role": "agent", "text": "".join(text_parts)})
        if audio_chunk_count:
            await websocket.send_json({"type": "status", "message": f"已收到 Omni 音频：{audio_chunk_count} 段，约 {audio_byte_count} 字节。"})
            await websocket.send_json({"type": "audio_done"})
        else:
            await websocket.send_json({"type": "status", "message": "Qwen-Omni 本轮没有返回音频。"})
    except httpx.HTTPError as exc:
        await websocket.send_json({"type": "error", "message": f"无法连接 Qwen-Omni API：{exc}"})


def find_report_by_interview_id(interview_id: str, user_id: str) -> dict | None:
    return one(
        """
        SELECT reports.id, reports.user_id, reports.interview_id,
               interviews.target_role, interviews.interview_type, interviews.status AS interview_status,
               reports.total_score, reports.grade, reports.pass_recommendation,
               reports.ability_radar, reports.agent_feedback, reports.timeline_review,
               reports.summary, reports.suggestions, reports.created_at, reports.updated_at
        FROM interview_reports AS reports
        JOIN interview_sessions AS interviews ON interviews.id = reports.interview_id
        WHERE reports.interview_id = ? AND reports.user_id = ?
        """,
        (interview_id, user_id),
    )


def find_report_by_id(report_id: str, user_id: str) -> dict | None:
    return one(
        """
        SELECT reports.id, reports.user_id, reports.interview_id,
               interviews.target_role, interviews.interview_type, interviews.status AS interview_status,
               reports.total_score, reports.grade, reports.pass_recommendation,
               reports.ability_radar, reports.agent_feedback, reports.timeline_review,
               reports.summary, reports.suggestions, reports.created_at, reports.updated_at
        FROM interview_reports AS reports
        JOIN interview_sessions AS interviews ON interviews.id = reports.interview_id
        WHERE reports.id = ? AND reports.user_id = ?
        """,
        (report_id, user_id),
    )


def list_reports_by_user_id(user_id: str) -> list[dict]:
    return all_rows(
        """
        SELECT reports.id, reports.user_id, reports.interview_id,
               interviews.target_role, interviews.interview_type, interviews.status AS interview_status,
               reports.total_score, reports.grade, reports.pass_recommendation,
               reports.summary, reports.created_at, reports.updated_at
        FROM interview_reports AS reports
        JOIN interview_sessions AS interviews ON interviews.id = reports.interview_id
        WHERE reports.user_id = ?
        ORDER BY reports.updated_at DESC, reports.created_at DESC
        """,
        (user_id,),
    )


def list_full_reports_by_user_id(user_id: str) -> list[dict]:
    return all_rows(
        """
        SELECT reports.id, reports.user_id, reports.interview_id,
               interviews.target_role, interviews.interview_type, interviews.status AS interview_status,
               reports.total_score, reports.ability_radar, reports.created_at, reports.updated_at
        FROM interview_reports AS reports
        JOIN interview_sessions AS interviews ON interviews.id = reports.interview_id
        WHERE reports.user_id = ?
        ORDER BY reports.updated_at ASC, reports.created_at ASC
        """,
        (user_id,),
    )


def compute_dimension_trend(values: list[int]) -> str:
    if len(values) < 2:
        return "stable"
    delta = values[-1] - values[0]
    if delta >= 5:
        return "up"
    if delta <= -5:
        return "down"
    return "stable"


def find_user_skill_stats(user_id: str) -> dict | None:
    return one(
        """
        SELECT id, user_id, total_interviews, completed_interviews, average_total_score,
               technical_depth_avg, expression_clarity_avg, business_understanding_avg,
               dimension_trends, weak_points, recent_training_focus, updated_at
        FROM user_skill_stats
        WHERE user_id = ?
        """,
        (user_id,),
    )


def refresh_user_skill_stats(user_id: str) -> dict:
    interview_counts = one(
        """
        SELECT
          COUNT(*) AS total_interviews,
          SUM(CASE WHEN status = 'completed' THEN 1 ELSE 0 END) AS completed_interviews
        FROM interview_sessions
        WHERE user_id = ?
        """,
        (user_id,),
    ) or {"total_interviews": 0, "completed_interviews": 0}
    reports = list_full_reports_by_user_id(user_id)
    dimensions = ["technical_depth", "expression_clarity", "business_understanding"]
    dimension_values: dict[str, list[int]] = {dimension: [] for dimension in dimensions}
    for report in reports:
        radar = parse_json_object(report.get("ability_radar"))
        for dimension in dimensions:
            value = radar.get(dimension)
            if isinstance(value, (int, float)):
                dimension_values[dimension].append(round(value))

    dimension_averages = {dimension: clamp_stat_score(average(values, 0)) for dimension, values in dimension_values.items()}
    dimension_trends = {dimension: compute_dimension_trend(values) for dimension, values in dimension_values.items()}
    weak_points = sorted(
        [{"dimension": dimension, "score": score} for dimension, score in dimension_averages.items() if score > 0],
        key=lambda item: item["score"],
    )[:2]
    recent_focus = reports[-1]["target_role"] if reports else None
    average_total_score = clamp_stat_score(average([report["total_score"] for report in reports], 0))
    existing = one("SELECT id FROM user_skill_stats WHERE user_id = ?", (user_id,))
    stats_id = existing["id"] if existing else str(uuid4())

    if existing:
        db.execute(
            """
            UPDATE user_skill_stats
            SET total_interviews = ?, completed_interviews = ?, average_total_score = ?,
                technical_depth_avg = ?, expression_clarity_avg = ?, business_understanding_avg = ?,
                dimension_trends = ?, weak_points = ?, recent_training_focus = ?,
                updated_at = CURRENT_TIMESTAMP
            WHERE user_id = ?
            """,
            (
                interview_counts.get("total_interviews") or 0,
                interview_counts.get("completed_interviews") or 0,
                average_total_score,
                dimension_averages["technical_depth"],
                dimension_averages["expression_clarity"],
                dimension_averages["business_understanding"],
                json.dumps(dimension_trends, ensure_ascii=False),
                json.dumps(weak_points, ensure_ascii=False),
                recent_focus,
                user_id,
            ),
        )
    else:
        db.execute(
            """
            INSERT INTO user_skill_stats (
              id, user_id, total_interviews, completed_interviews, average_total_score,
              technical_depth_avg, expression_clarity_avg, business_understanding_avg,
              dimension_trends, weak_points, recent_training_focus
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                stats_id,
                user_id,
                interview_counts.get("total_interviews") or 0,
                interview_counts.get("completed_interviews") or 0,
                average_total_score,
                dimension_averages["technical_depth"],
                dimension_averages["expression_clarity"],
                dimension_averages["business_understanding"],
                json.dumps(dimension_trends, ensure_ascii=False),
                json.dumps(weak_points, ensure_ascii=False),
                recent_focus,
            ),
        )
    db.commit()
    return find_user_skill_stats(user_id)


def build_user_dimension_stats(stats: dict) -> list[dict]:
    trends = parse_json_object(stats.get("dimension_trends"))
    return [
        {
            "key": "technical_depth",
            "label": "技术深度",
            "average_score": stats["technical_depth_avg"],
            "trend": trends.get("technical_depth", "stable"),
        },
        {
            "key": "expression_clarity",
            "label": "表达清晰度",
            "average_score": stats["expression_clarity_avg"],
            "trend": trends.get("expression_clarity", "stable"),
        },
        {
            "key": "business_understanding",
            "label": "业务理解",
            "average_score": stats["business_understanding_avg"],
            "trend": trends.get("business_understanding", "stable"),
        },
    ]


@app.get("/api/health")
def health():
    return {"ok": True, "databasePath": get_database_path()}


@app.get("/api/auth/me")
def auth_me(request: Request):
    return {"user": sanitize_user(find_user_by_session(request))}


@app.post("/api/auth/register", status_code=201)
async def auth_register(request: Request, response: Response, body: dict | None = None):
    body = json_body(body)
    email = normalize_email(body.get("email"))
    password = str(body.get("password") or "")
    name = str(body.get("name") or "").strip() or (email.split("@")[0] if "@" in email else "新用户")
    if not is_valid_email(email):
        raise error(400, "请输入有效邮箱。")
    if len(password) < 8:
        raise error(400, "密码至少需要 8 位。")
    if len(name) > 60:
        raise error(400, "昵称不能超过 60 个字符。")
    if one("SELECT id FROM users WHERE email = ?", (email,)):
        raise error(409, "这个邮箱已经注册。")

    user_id = str(uuid4())
    with db:
        db.execute("INSERT INTO users (id, email, password_hash, name) VALUES (?, ?, ?, ?)", (user_id, email, hash_password(password), name))
        db.execute("INSERT INTO profiles (id, user_id, nickname) VALUES (?, ?, ?)", (str(uuid4()), user_id, name))
    create_session(response, user_id, request.headers.get("user-agent"))
    user = one("SELECT id, email, name, status, created_at, last_login_at FROM users WHERE id = ?", (user_id,))
    return {"user": sanitize_user(user)}


@app.post("/api/auth/login")
async def auth_login(request: Request, response: Response, body: dict | None = None):
    body = json_body(body)
    email = normalize_email(body.get("email"))
    password = str(body.get("password") or "")
    if not is_valid_email(email) or not password:
        raise error(400, "请输入邮箱和密码。")
    if is_login_limited(request, email):
        raise error(429, "登录尝试过于频繁，请稍后再试。")

    user = one("SELECT id, email, password_hash, name, status, created_at, last_login_at FROM users WHERE email = ?", (email,))
    if not user or user["status"] != "normal" or not verify_password(password, user["password_hash"]):
        record_failed_login(request, email)
        raise error(401, "邮箱或密码不正确。")

    clear_failed_logins(request, email)
    db.execute("UPDATE users SET last_login_at = CURRENT_TIMESTAMP, updated_at = CURRENT_TIMESTAMP WHERE id = ?", (user["id"],))
    db.commit()
    create_session(response, user["id"], request.headers.get("user-agent"))
    current_user = one("SELECT id, email, name, status, created_at, last_login_at FROM users WHERE id = ?", (user["id"],))
    return {"user": sanitize_user(current_user)}


@app.post("/api/auth/logout")
def auth_logout(request: Request, response: Response):
    token = request.cookies.get(SESSION_COOKIE_NAME)
    if token:
        db.execute("DELETE FROM sessions WHERE token_hash = ?", (hash_token(token),))
        db.commit()
    clear_session(response)
    return {"ok": True}


@app.post("/api/auth/password-reset/request")
def password_reset_request(body: dict | None = None):
    body = json_body(body)
    email = normalize_email(body.get("email"))
    user = one("SELECT id FROM users WHERE email = ? AND status = ?", (email, "normal")) if is_valid_email(email) else None
    dev_reset_token = None
    if user:
        token = create_token()
        db.execute(
            "INSERT INTO password_reset_tokens (id, user_id, token_hash, expires_at) VALUES (?, ?, ?, ?)",
            (str(uuid4()), user["id"], hash_token(token), now_iso_after(RESET_MAX_AGE_SECONDS)),
        )
        db.commit()
        if os.environ.get("APP_ENV") != "production":
            dev_reset_token = token
    return {"ok": True, "message": "如果邮箱存在，我们会发送密码重置链接。", "devResetToken": dev_reset_token}


@app.get("/api/profile")
def get_profile(user: dict = Depends(require_auth)):
    ensure_profile(user)
    return {"profile": find_profile_by_user_id(user["id"])}


@app.put("/api/profile")
def update_profile(body: dict | None = None, user: dict = Depends(require_auth)):
    profile = parse_profile_input(json_body(body), user["name"])
    ensure_profile(user)
    db.execute(
        """
        UPDATE profiles
        SET nickname = ?, avatar_url = ?, target_role = ?, experience_level = ?, company_type = ?,
            target_city = ?, expected_salary = ?, years_of_experience = ?, education_level = ?,
            skills = ?, project_keywords = ?, resume_text = ?, project_experience = ?,
            portfolio_links = ?, preferred_interview_type = ?, preferred_difficulty = ?,
            preferred_interviewer_style = ?, updated_at = CURRENT_TIMESTAMP
        WHERE user_id = ?
        """,
        (
            profile["nickname"],
            profile["avatar_url"],
            profile["target_role"],
            profile["experience_level"],
            profile["company_type"],
            profile["target_city"],
            profile["expected_salary"],
            profile["years_of_experience"],
            profile["education_level"],
            profile["skills"],
            profile["project_keywords"],
            profile["resume_text"],
            profile["project_experience"],
            profile["portfolio_links"],
            profile["preferred_interview_type"],
            profile["preferred_difficulty"],
            profile["preferred_interviewer_style"],
            user["id"],
        ),
    )
    db.execute("UPDATE users SET name = ?, updated_at = CURRENT_TIMESTAMP WHERE id = ?", (profile["nickname"], user["id"]))
    db.commit()
    return {"profile": find_profile_by_user_id(user["id"])}


def extract_text_from_resume_file(filename: str, content: bytes) -> tuple[str, str | None]:
    """Extract plain text from an uploaded resume file."""
    extension = (filename.rsplit(".", 1)[-1] if "." in filename else "").lower()

    if extension in ("txt", "md", "json"):
        return content.decode("utf-8", errors="replace").strip(), None

    if extension == "pdf":
        try:
            from pypdf import PdfReader
        except ImportError:
            return "", "PDF 解析功能未启用（缺少 pypdf），请复制文字粘贴到文本框。"

        try:
            reader = PdfReader(io.BytesIO(content))
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
            return text.strip(), None
        except Exception as exc:
            return "", f"PDF 解析失败：{exc}"

    if extension == "docx":
        try:
            from docx import Document
        except ImportError:
            return "", "Word 文档解析功能未启用（缺少 python-docx），请复制文字粘贴到文本框。"

        try:
            doc = Document(io.BytesIO(content))
            lines = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        lines.append(" | ".join(cells))
            text = "\n".join(lines)
            return text.strip(), None
        except Exception as exc:
            return "", f"Word 文档解析失败：{exc}"

    if extension == "doc":
        return "", "暂不支持 .doc 旧格式，请另存为 .docx 后重新上传。"

    if extension in ("png", "jpg", "jpeg"):
        try:
            import pytesseract
            from PIL import Image
        except ImportError:
            return "", "图片 OCR 功能未启用（缺少 pytesseract 或 Pillow），请复制文字粘贴到文本框。"

        try:
            image = Image.open(io.BytesIO(content))
            try:
                text = pytesseract.image_to_string(image, lang="chi_sim+eng")
            except pytesseract.pytesseract.TesseractError as exc:
                if "Error opening data file" not in str(exc) and "Failed loading language" not in str(exc):
                    raise
                text = pytesseract.image_to_string(image, lang="eng")
            return text.strip(), None
        except pytesseract.pytesseract.TesseractNotFoundError:
            return "", "图片 OCR 功能未启用（缺少 Tesseract 引擎），请复制文字粘贴到文本框。"
        except pytesseract.pytesseract.TesseractError as exc:
            return "", f"图片 OCR 语言包不可用或识别失败：{exc}"
        except Exception as exc:
            return "", f"图片识别失败：{exc}"

    if extension:
        return "", f"不支持的文件格式：.{extension}"
    return "", "不支持的文件格式，请上传 txt、md、json、PDF、Word 文档或图片。"


@app.post("/api/profile/resume-upload")
async def upload_resume_file(
    file: UploadFile = File(...),
    user: dict = Depends(require_auth),
):
    content = await file.read()
    if len(content) > MAX_RESUME_FILE_BYTES:
        raise HTTPException(status_code=413, detail="文件过大，请上传 10MB 以内的简历文件。")

    text, error_message = extract_text_from_resume_file(file.filename or "", content)
    if error_message:
        raise HTTPException(status_code=422, detail=error_message)

    if not text.strip():
        raise HTTPException(status_code=422, detail="文件内容为空或未能提取到文本，请检查文件或直接粘贴简历文本。")

    saved_text = text[:12000]

    ensure_profile(user)
    db.execute(
        "UPDATE profiles SET resume_text = ?, updated_at = CURRENT_TIMESTAMP WHERE user_id = ?",
        (saved_text, user["id"]),
    )
    db.commit()

    return {
        "filename": file.filename,
        "text": saved_text,
        "char_count": len(saved_text),
        "original_char_count": len(text),
        "truncated": len(text) > len(saved_text),
        "profile": find_profile_by_user_id(user["id"]),
    }


@app.get("/api/profile/resume-analysis")
def get_resume_analysis(user: dict = Depends(require_auth)):
    ensure_profile(user)
    profile = find_profile_by_user_id(user["id"])
    current_hash = resume_source_hash(profile)
    row = find_resume_analysis_by_user_id(user["id"])
    serialized = serialize_resume_analysis(row)
    return {
        "resume_analysis": serialized,
        "current_source_hash": current_hash,
        "stale": bool(serialized and serialized.get("source_hash") != current_hash),
    }


@app.post("/api/profile/resume-analysis")
def create_resume_analysis(body: dict | None = None, user: dict = Depends(require_auth)):
    force = bool(json_body(body).get("force"))
    return {"resume_analysis": ensure_resume_analysis_for_user(user, force=force)}


@app.post("/api/interviews", status_code=201)
def create_interview(body: dict | None = None, user: dict = Depends(require_auth)):
    interview = parse_interview_input(json_body(body))
    interview_id = str(uuid4())
    with db:
        db.execute(
            """
            INSERT INTO interview_sessions (
              id, user_id, target_role, experience_level, interview_type, company_context,
              focus_areas, difficulty, interviewer_style, status
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                interview_id,
                user["id"],
                interview["target_role"],
                interview["experience_level"],
                interview["interview_type"],
                interview["company_context"],
                interview["focus_areas"],
                interview["difficulty"],
                interview["interviewer_style"],
                "draft",
            ),
        )
        create_default_agents(interview_id)
    return {"interview": find_interview_by_user_id(interview_id, user["id"]), "agents": list_agents_by_interview_id(interview_id)}


@app.get("/api/interviews")
def list_interviews(request: Request, user: dict = Depends(require_auth)):
    raw_status = request.query_params.get("status")
    status = normalize_status(raw_status) if raw_status else None
    if raw_status and not status:
        raise error(400, "面试状态不正确。")
    if status:
        interviews = all_rows(
            """
            SELECT id, user_id, target_role, experience_level, interview_type, company_context,
                   focus_areas, difficulty, interviewer_style, status, created_at, updated_at,
                   started_at, completed_at
            FROM interview_sessions
            WHERE user_id = ? AND status = ?
            ORDER BY updated_at DESC, created_at DESC
            """,
            (user["id"], status),
        )
    else:
        interviews = all_rows(
            """
            SELECT id, user_id, target_role, experience_level, interview_type, company_context,
                   focus_areas, difficulty, interviewer_style, status, created_at, updated_at,
                   started_at, completed_at
            FROM interview_sessions
            WHERE user_id = ?
            ORDER BY updated_at DESC, created_at DESC
            """,
            (user["id"],),
        )
    return {"interviews": interviews}


@app.get("/api/interviews/{interview_id}")
def get_interview(interview_id: str, user: dict = Depends(require_auth)):
    interview = find_interview_by_user_id(interview_id, user["id"])
    if not interview:
        raise error(404, "面试不存在。")
    return {"interview": interview}


@app.patch("/api/interviews/{interview_id}")
def update_interview(interview_id: str, body: dict | None = None, user: dict = Depends(require_auth)):
    interview = find_interview_by_user_id(interview_id, user["id"])
    if not interview:
        raise error(404, "面试不存在。")
    if interview["status"] == "completed":
        raise error(409, "已完成的面试不能继续修改。")
    parsed = parse_interview_input(json_body(body), partial=True)
    next_interview = {**interview, **parsed}
    db.execute(
        """
        UPDATE interview_sessions
        SET target_role = ?, experience_level = ?, interview_type = ?, company_context = ?,
            focus_areas = ?, difficulty = ?, interviewer_style = ?, updated_at = CURRENT_TIMESTAMP
        WHERE id = ? AND user_id = ?
        """,
        (
            next_interview["target_role"],
            next_interview["experience_level"],
            next_interview["interview_type"],
            next_interview["company_context"],
            next_interview["focus_areas"],
            next_interview["difficulty"],
            next_interview["interviewer_style"],
            interview["id"],
            user["id"],
        ),
    )
    db.commit()
    return {"interview": find_interview_by_user_id(interview["id"], user["id"])}


@app.post("/api/interviews/{interview_id}/start")
def start_interview(interview_id: str, user: dict = Depends(require_auth)):
    interview = find_interview_by_user_id(interview_id, user["id"])
    if not interview:
        raise error(404, "面试不存在。")
    if interview["status"] == "completed":
        raise error(409, "已完成的面试不能重新开始。")
    if interview["status"] == "draft":
        db.execute(
            """
            UPDATE interview_sessions
            SET status = 'running', started_at = COALESCE(started_at, CURRENT_TIMESTAMP),
                updated_at = CURRENT_TIMESTAMP
            WHERE id = ? AND user_id = ?
            """,
            (interview["id"], user["id"]),
        )
        db.commit()
    return {"interview": find_interview_by_user_id(interview["id"], user["id"])}


@app.post("/api/interviews/{interview_id}/opening-question")
def create_opening_question(interview_id: str, body: dict | None = None, user: dict = Depends(require_auth)):
    interview = find_interview_by_user_id(interview_id, user["id"])
    if not interview:
        raise error(404, "面试不存在。")
    if interview["status"] == "completed":
        raise error(409, "已完成的面试不能继续生成首问。")
    messages = list_messages_by_interview_id(interview["id"])
    if messages and not bool(json_body(body).get("force")):
        raise error(409, "当前面试已有消息，不能重复生成首问。")

    ensure_default_agents(interview["id"])
    agents = list_agents_by_interview_id(interview["id"])
    active_agent = next((agent for agent in agents if agent["status"] == "active"), None)
    if not active_agent:
        active_agent = next((agent for agent in agents if agent["status"] != "completed"), agents[0] if agents else None)

    resume_analysis_row = ensure_resume_analysis_for_user(user)
    analysis = resume_analysis_row.get("analysis") or {}
    selected_resume_point = select_resume_point(analysis, interview, messages)
    result = generate_opening_question(
        interview=interview,
        active_agent=active_agent,
        resume_analysis=analysis,
        selected_resume_point=selected_resume_point,
    )
    return {
        "question": result["question"],
        "provider": result.get("provider"),
        "fallback": bool(result.get("fallback")),
        "error": result.get("error"),
        "selected_resume_point": selected_resume_point,
        "resume_analysis": resume_analysis_row,
    }


@app.post("/api/interviews/{interview_id}/finish")
def finish_interview(interview_id: str, user: dict = Depends(require_auth)):
    interview = find_interview_by_user_id(interview_id, user["id"])
    if not interview:
        raise error(404, "面试不存在。")
    if interview["status"] == "completed":
        return {"interview": interview}
    db.execute(
        """
        UPDATE interview_sessions
        SET status = 'completed', started_at = COALESCE(started_at, CURRENT_TIMESTAMP),
            completed_at = COALESCE(completed_at, CURRENT_TIMESTAMP), updated_at = CURRENT_TIMESTAMP
        WHERE id = ? AND user_id = ?
        """,
        (interview["id"], user["id"]),
    )
    db.commit()
    return {"interview": find_interview_by_user_id(interview["id"], user["id"])}


@app.get("/api/interviews/{interview_id}/agents")
def list_agents(interview_id: str, user: dict = Depends(require_auth)):
    interview = find_interview_by_user_id(interview_id, user["id"])
    if not interview:
        raise error(404, "面试不存在。")
    ensure_default_agents(interview["id"])
    return {"agents": list_agents_by_interview_id(interview["id"])}


@app.post("/api/interviews/{interview_id}/agents", status_code=201)
def create_agent(interview_id: str, body: dict | None = None, user: dict = Depends(require_auth)):
    interview = find_interview_by_user_id(interview_id, user["id"])
    if not interview:
        raise error(404, "面试不存在。")
    if interview["status"] == "completed":
        raise error(409, "已完成的面试不能新增 Agent。")
    agent = parse_agent_input(json_body(body))
    agent_id = str(uuid4())
    order_index = agent["order_index"] if isinstance(agent.get("order_index"), int) else next_agent_order_index(interview["id"])
    db.execute(
        """
        INSERT INTO interview_agents (
          id, interview_id, agent_name, agent_type, agent_role, strategy, order_index, status
        )
        VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            agent_id,
            interview["id"],
            agent["agent_name"],
            agent["agent_type"],
            agent["agent_role"],
            agent.get("strategy"),
            order_index,
            agent.get("status") or "pending",
        ),
    )
    db.commit()
    return {"agent": find_agent_by_interview_id(agent_id, interview["id"])}


@app.patch("/api/interviews/{interview_id}/agents/{agent_id}")
def update_agent(interview_id: str, agent_id: str, body: dict | None = None, user: dict = Depends(require_auth)):
    interview = find_interview_by_user_id(interview_id, user["id"])
    if not interview:
        raise error(404, "面试不存在。")
    if interview["status"] == "completed":
        raise error(409, "已完成的面试不能修改 Agent。")
    existing = find_agent_by_interview_id(agent_id, interview["id"])
    if not existing:
        raise error(404, "Agent 不存在。")
    parsed = parse_agent_input(json_body(body), partial=True)
    next_agent = {**existing, **parsed}
    db.execute(
        """
        UPDATE interview_agents
        SET agent_name = ?, agent_type = ?, agent_role = ?, strategy = ?, order_index = ?,
            status = ?, updated_at = CURRENT_TIMESTAMP
        WHERE id = ? AND interview_id = ?
        """,
        (
            next_agent["agent_name"],
            next_agent["agent_type"],
            next_agent["agent_role"],
            next_agent.get("strategy"),
            next_agent["order_index"],
            next_agent["status"],
            existing["id"],
            interview["id"],
        ),
    )
    db.commit()
    return {"agent": find_agent_by_interview_id(existing["id"], interview["id"])}


@app.get("/api/interviews/{interview_id}/messages")
def list_messages(interview_id: str, user: dict = Depends(require_auth)):
    interview = find_interview_by_user_id(interview_id, user["id"])
    if not interview:
        owner = one(
            "SELECT id, user_id, status, created_at, updated_at FROM interview_sessions WHERE id = ?",
            (interview_id,),
        )
        reason = "missing_interview"
        if owner:
            reason = "owner_matches_current_user" if owner.get("user_id") == user["id"] else "owner_belongs_to_other_user"
        print(
            "Messages 404 debug:",
            json.dumps(
                {
                    "interview_id": interview_id,
                    "current_user_id": user["id"],
                    "owner": owner,
                    "reason": reason,
                },
                ensure_ascii=False,
            ),
            flush=True,
        )
        if owner and owner.get("user_id") == user["id"]:
            print(
                "Messages 404 recovered:",
                json.dumps({"interview_id": interview_id, "current_user_id": user["id"]}, ensure_ascii=False),
                flush=True,
            )
            return {"messages": list_messages_by_interview_id(owner["id"]), "recovered": True}
        raise error(404, "面试不存在。")
    return {"messages": list_messages_by_interview_id(interview["id"])}


@app.post("/api/interviews/{interview_id}/messages", status_code=201)
def create_message(interview_id: str, body: dict | None = None, user: dict = Depends(require_auth)):
    interview = find_interview_by_user_id(interview_id, user["id"])
    if not interview:
        raise error(404, "面试不存在。")
    if interview["status"] == "completed":
        raise error(409, "已完成的面试不能继续写入消息。")
    message = parse_message_input(json_body(body))
    if message["agent_id"] and not find_agent_by_interview_id(message["agent_id"], interview["id"]):
        raise error(400, "Agent 不属于当前面试。")
    message_id = str(uuid4())
    db.execute(
        """
        INSERT INTO interview_messages (
          id, interview_id, agent_id, sender_type, message_type, content, transcript_text, order_index
        )
        VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            message_id,
            interview["id"],
            message["agent_id"],
            message["sender_type"],
            message["message_type"],
            message["content"],
            message["transcript_text"],
            next_message_order_index(interview["id"]),
        ),
    )
    db.execute("UPDATE interview_sessions SET updated_at = CURRENT_TIMESTAMP WHERE id = ? AND user_id = ?", (interview["id"], user["id"]))
    db.commit()
    return {"message": find_message_by_interview_id(message_id, interview["id"])}


@app.post("/api/interviews/{interview_id}/follow-up")
def generate_follow_up(interview_id: str, body: dict | None = None, user: dict = Depends(require_auth)):
    interview = find_interview_by_user_id(interview_id, user["id"])
    if not interview:
        raise error(404, "面试不存在。")
    if interview["status"] == "completed":
        raise error(409, "已完成的面试不能继续生成追问。")
    last_answer, message = read_text_field(json_body(body), "last_answer", "候选人回答", 20000, True)
    if message:
        raise error(400, message)

    ensure_default_agents(interview["id"])
    agents = list_agents_by_interview_id(interview["id"])
    active_agent = next((agent for agent in agents if agent["status"] == "active"), None)
    if not active_agent:
        active_agent = next((agent for agent in agents if agent["status"] != "completed"), agents[0] if agents else None)
    messages = list_messages_by_interview_id(interview["id"])
    resume_analysis = None
    ensure_profile(user)
    profile = find_profile_by_user_id(user["id"])
    if resume_source_text(profile):
        resume_analysis = ensure_resume_analysis_for_user(user).get("analysis")
    try:
        follow_up = generate_kimi_followup(
            interview=interview,
            active_agent=active_agent,
            messages=messages,
            last_answer=last_answer,
            resume_analysis=resume_analysis,
        )
    except KimiFollowupError as exc:
        raise error(502, str(exc))
    return follow_up


@app.post("/api/interviews/{interview_id}/realtime/sdp", response_class=PlainTextResponse)
async def realtime_sdp(interview_id: str, request: Request, user: dict = Depends(require_auth)):
    interview = find_interview_by_user_id(interview_id, user["id"])
    if not interview:
        raise error(404, "面试不存在。")
    if interview["status"] == "completed":
        raise error(409, "已完成的面试不能开启语音通话。")
    api_key = os.environ.get("OPENAI_API_KEY") or os.environ.get("AI_API_KEY")
    if not api_key or api_key == "your-api-key-here":
        raise error(500, "服务端未配置 OPENAI_API_KEY，无法开启实时语音。")
    sdp = (await request.body()).decode("utf-8")
    if not sdp.strip():
        raise error(400, "缺少 WebRTC SDP offer。")
    print(
        "Realtime SDP received:",
        {"length": len(sdp), "first_line": sdp.splitlines()[0] if sdp.splitlines() else ""},
        flush=True,
    )

    agents = list_agents_by_interview_id(interview["id"])
    messages = list_messages_by_interview_id(interview["id"])
    files = {
        "sdp": (None, sdp),
        "session": (None, json.dumps(build_openai_realtime_session_config(interview, agents, messages), ensure_ascii=False)),
    }
    try:
        async with httpx.AsyncClient(timeout=30) as client:
            response = await client.post(
                "https://api.openai.com/v1/realtime/calls",
                headers={"Authorization": f"Bearer {api_key}"},
                files=files,
            )
    except httpx.HTTPError as exc:
        raise error(502, f"无法连接 OpenAI Realtime API：{exc}")

    if response.status_code >= 400:
        message = "创建实时语音会话失败。"
        response_body = response.text or ""
        try:
            data = response.json()
            message = data.get("error", {}).get("message") or data.get("message") or message
        except ValueError:
            message = response_body or message
        print(
            "OpenAI Realtime calls failed:",
            response.status_code,
            response_body[:2000],
            flush=True,
        )
        raise error(response.status_code, message)
    return PlainTextResponse(response.text, media_type="application/sdp")


@app.post("/api/interviews/{interview_id}/qwen/tts")
def qwen_tts(interview_id: str, body: dict | None = None, user: dict = Depends(require_auth)):
    interview = find_interview_by_user_id(interview_id, user["id"])
    if not interview:
        raise error(404, "面试不存在。")
    text, message = read_text_field(json_body(body), "text", "语音合成文本", 20000, True)
    if message:
        raise error(400, message)
    return build_qwen_tts_response(text, provider_preference="standard")


@app.post("/api/interviews/{interview_id}/qwen/realtime-tts")
def qwen_realtime_tts(interview_id: str, body: dict | None = None, user: dict = Depends(require_auth)):
    interview = find_interview_by_user_id(interview_id, user["id"])
    if not interview:
        raise error(404, "面试不存在。")
    text, message = read_text_field(json_body(body), "text", "实时语音合成文本", 20000, True)
    if message:
        raise error(400, message)
    return build_qwen_tts_response(text, provider_preference="realtime")


@app.post("/api/interviews/{interview_id}/qwen/speech")
def qwen_speech(interview_id: str, body: dict | None = None, user: dict = Depends(require_auth)):
    interview = find_interview_by_user_id(interview_id, user["id"])
    if not interview:
        raise error(404, "面试不存在。")
    payload = json_body(body)
    text, message = read_text_field(payload, "text", "语音合成文本", 20000, True)
    if message:
        raise error(400, message)
    provider = str(payload.get("provider") or "auto")
    return build_qwen_tts_response(text, provider_preference=provider)


@app.post("/api/interviews/{interview_id}/qwen/omni-realtime/sdp")
async def qwen_omni_realtime_sdp(interview_id: str, body: dict | None = None, user: dict = Depends(require_auth)):
    interview = find_interview_by_user_id(interview_id, user["id"])
    if not interview:
        raise error(404, "面试不存在。")
    if interview["status"] == "completed":
        raise error(409, "已完成的面试不能开启 Qwen-Omni Realtime。")

    payload = json_body(body)
    sdp = str(payload.get("sdp") or "")
    offer_type = str(payload.get("type") or "offer")
    if not sdp.strip() or offer_type != "offer":
        raise error(400, "请提供有效的 WebRTC offer。")

    api_key = _qwen_realtime_api_key()
    if not api_key or api_key in {"your-dashscope-api-key", "your-api-key-here"}:
        raise error(500, "服务端未配置 DASHSCOPE_API_KEY 或 QWEN_OMNI_REALTIME_API_KEY，无法开启 Qwen-Omni Realtime。")

    agents = list_agents_by_interview_id(interview["id"])
    messages = list_messages_by_interview_id(interview["id"])
    session_config = build_realtime_session_config(interview, agents, messages)
    session_config["provider"] = "qwen-omni-realtime-webrtc"
    session_config["interview_id"] = interview_id
    start_payload = {
        "currentQuestion": str(payload.get("currentQuestion") or ""),
        "activeAgent": str(payload.get("activeAgent") or ""),
    }

    try:
        model = _qwen_realtime_model()
        upstream_url = _qwen_realtime_webrtc_url(model)
        session_update = build_qwen_realtime_session_update(session_config, start_payload)
    except QwenRealtimeTtsError as exc:
        raise error(500, str(exc))

    try:
        async with httpx.AsyncClient(timeout=float(os.environ.get("QWEN_OMNI_REALTIME_WEBRTC_TIMEOUT", "30"))) as client:
            response = await client.post(
                upstream_url,
                headers={
                    "Authorization": f"Bearer {api_key}",
                    "Content-Type": "application/sdp",
                    "Accept": "application/sdp",
                },
                content=sdp,
            )
    except httpx.HTTPError as exc:
        raise error(502, f"无法连接 Qwen-Omni Realtime WebRTC：{exc}")

    if response.status_code >= 400:
        response_body = response.text or ""
        message = response_body or "创建 Qwen-Omni Realtime WebRTC 会话失败。"
        try:
            data = response.json()
            message = data.get("error", {}).get("message") or data.get("message") or message
        except ValueError:
            pass
        print("Qwen-Omni WebRTC SDP failed:", response.status_code, response_body[:2000], flush=True)
        raise error(response.status_code, message[:1200])

    answer_sdp = response.text.strip()
    if not answer_sdp:
        raise error(502, "Qwen-Omni Realtime 未返回有效 SDP answer。")

    return {
        "provider": "qwen-omni-realtime-webrtc",
        "model": model,
        "answer": {"type": "answer", "sdp": answer_sdp},
        "session_update": session_update,
    }


@app.post("/api/interviews/{interview_id}/qwen/omni-webrtc/offer")
async def qwen_omni_webrtc_offer(interview_id: str, body: dict | None = None, user: dict = Depends(require_auth)):
    return await qwen_omni_realtime_sdp(interview_id, body, user)


@app.post("/api/interviews/{interview_id}/qwen/omni-webrtc/stop")
async def qwen_omni_webrtc_stop(interview_id: str, _body: dict | None = None, user: dict = Depends(require_auth)):
    interview = find_interview_by_user_id(interview_id, user["id"])
    if not interview:
        raise error(404, "面试不存在。")
    return {
        "events": [
            {
                "type": "status",
                "message": "Qwen-Omni WebRTC 已改为浏览器直连链路，断开连接不再需要服务端提交音频。",
            }
        ]
    }


@app.websocket("/ws/interviews/{interview_id}/qwen/omni-realtime")
async def qwen_omni_realtime_ws(websocket: WebSocket, interview_id: str):
    await websocket.accept()

    user = find_user_by_session_token(websocket.cookies.get(SESSION_COOKIE_NAME))
    if not user:
        await websocket.send_json({"type": "error", "message": "请先登录后再连接千问 Omni 实时通话。"})
        await websocket.close(code=1008)
        return

    interview = find_interview_by_user_id(interview_id, user["id"])
    if not interview:
        await websocket.send_json({"type": "error", "message": "面试不存在。"})
        await websocket.close(code=1008)
        return
    if interview["status"] == "completed":
        await websocket.send_json({"type": "error", "message": "已完成的面试不能开启千问 Omni 实时通话。"})
        await websocket.close(code=1008)
        return

    agents = list_agents_by_interview_id(interview["id"])
    messages = list_messages_by_interview_id(interview["id"])
    session_config = build_realtime_session_config(interview, agents, messages)
    session_config["provider"] = "qwen-omni-self-hosted"
    session_config["interview_id"] = interview_id

    upstream_url = os.environ.get("QWEN_OMNI_REALTIME_WS_URL", "").strip()
    if not upstream_url:
        await websocket.send_json({"type": "status", "message": "未配置自部署 WS 网关，改用 Qwen-Omni 官方 HTTP API 试验模式。"})
        audio_chunks: list[bytes] = []
        start_payload: dict = {}
        mime_type = "audio/webm"
        try:
            while True:
                message = await websocket.receive()
                if message.get("type") == "websocket.disconnect":
                    return
                if message.get("bytes") is not None:
                    audio_chunks.append(message["bytes"])
                    continue
                if message.get("text") is None:
                    continue
                try:
                    event = json.loads(message["text"])
                except ValueError:
                    continue
                event_type = event.get("type")
                if event_type == "start":
                    start_payload = event
                    mime_type = str(event.get("mimeType") or mime_type)
                    await websocket.send_json({"type": "status", "message": "正在录音，停止后提交给 Qwen-Omni API。"})
                if event_type == "stop":
                    await websocket.send_json({"type": "status", "message": "录音已收到，正在请求 Qwen-Omni。"})
                    latest_interview = find_interview_by_user_id(interview_id, user["id"]) or interview
                    latest_agents = list_agents_by_interview_id(interview["id"])
                    latest_messages = list_messages_by_interview_id(interview["id"])
                    round_session_config = build_realtime_session_config(latest_interview, latest_agents, latest_messages)
                    round_session_config["provider"] = "qwen-omni-http"
                    round_session_config["interview_id"] = interview_id
                    await _stream_qwen_omni_http_response(websocket, b"".join(audio_chunks), mime_type, round_session_config, start_payload)
                    audio_chunks = []
                    await websocket.send_json({"type": "status", "message": "本轮完成，可再次点击麦克风录下一轮。"})
        except WebSocketDisconnect:
            return
        except Exception as exc:
            await websocket.send_json({"type": "error", "message": f"Qwen-Omni HTTP 试验模式失败：{exc}"})
            await websocket.close(code=1011)
        return

    try:
        import asyncio
        import websockets
    except ImportError:
        await websocket.send_json({"type": "error", "message": "服务端缺少 websockets 依赖，无法桥接自部署 Qwen-Omni。"})
        await websocket.close(code=1011)
        return

    headers = {}
    upstream_key = os.environ.get("QWEN_OMNI_REALTIME_API_KEY", "").strip()
    if upstream_key:
        headers["Authorization"] = f"Bearer {upstream_key}"

    try:
        async with websockets.connect(
            upstream_url,
            additional_headers=headers or None,
            max_size=int(os.environ.get("QWEN_OMNI_REALTIME_MAX_MESSAGE_SIZE", "16777216")),
        ) as upstream:
            await upstream.send(json.dumps({"type": "session.start", "session": session_config}, ensure_ascii=False))
            await websocket.send_json({"type": "status", "message": "已连接自部署 Qwen-Omni 网关。"})

            async def forward_browser_to_upstream() -> None:
                while True:
                    message = await websocket.receive()
                    if message.get("type") == "websocket.disconnect":
                        break
                    if message.get("bytes") is not None:
                        await upstream.send(message["bytes"])
                    elif message.get("text") is not None:
                        await upstream.send(message["text"])

            async def forward_upstream_to_browser() -> None:
                async for message in upstream:
                    if isinstance(message, bytes):
                        await websocket.send_bytes(message)
                    else:
                        await websocket.send_text(message)

            done, pending = await asyncio.wait(
                {
                    asyncio.create_task(forward_browser_to_upstream()),
                    asyncio.create_task(forward_upstream_to_browser()),
                },
                return_when=asyncio.FIRST_COMPLETED,
            )
            for task in pending:
                task.cancel()
            for task in done:
                task.result()
    except WebSocketDisconnect:
        return
    except Exception as exc:
        await websocket.send_json({"type": "error", "message": f"千问 Omni 实时网关连接失败：{exc}"})
        await websocket.close(code=1011)


@app.post("/api/interviews/{interview_id}/evaluations", status_code=201)
def create_evaluation(
    interview_id: str,
    response: Response,
    body: dict | None = None,
    user: dict = Depends(require_auth),
):
    interview = find_interview_by_user_id(interview_id, user["id"])
    if not interview:
        raise error(404, "面试不存在。")
    message_id = str(json_body(body).get("message_id") or "").strip()
    if not message_id:
        raise error(400, "请提供要评价的 message_id。")
    message = find_message_by_interview_id(message_id, interview["id"])
    if not message:
        raise error(404, "消息不存在。")
    if message["sender_type"] != "candidate" or message["message_type"] not in {"answer", "transcript"}:
        raise error(400, "只能评价候选人的回答消息。")
    existing = find_evaluation_by_message_id(interview["id"], message["id"])
    if existing:
        if response:
            response.status_code = 200
        return {"evaluation": existing}
    evaluation_id = str(uuid4())
    evaluation = generate_mock_evaluation(message)
    db.execute(
        """
        INSERT INTO interview_evaluations (
          id, interview_id, message_id, agent_id, score, strengths, issues, suggestions, dimension_scores
        )
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            evaluation_id,
            interview["id"],
            message["id"],
            message["agent_id"],
            evaluation["score"],
            evaluation["strengths"],
            evaluation["issues"],
            evaluation["suggestions"],
            json.dumps(evaluation["dimension_scores"], ensure_ascii=False),
        ),
    )
    db.execute("UPDATE interview_sessions SET updated_at = CURRENT_TIMESTAMP WHERE id = ? AND user_id = ?", (interview["id"], user["id"]))
    db.commit()
    return {"evaluation": find_evaluation_by_id(interview["id"], evaluation_id)}


@app.get("/api/interviews/{interview_id}/evaluations")
def list_evaluations(interview_id: str, user: dict = Depends(require_auth)):
    interview = find_interview_by_user_id(interview_id, user["id"])
    if not interview:
        raise error(404, "面试不存在。")
    return {"evaluations": list_evaluations_by_interview_id(interview["id"])}


@app.post("/api/interviews/{interview_id}/report", status_code=201)
def create_report(interview_id: str, response: Response, user: dict = Depends(require_auth)):
    interview = find_interview_by_user_id(interview_id, user["id"])
    if not interview:
        raise error(404, "面试不存在。")
    if interview["status"] != "completed":
        raise error(409, "请先结束面试，再生成报告。")
    agents = list_agents_by_interview_id(interview["id"])
    messages = list_messages_by_interview_id(interview["id"])
    evaluations = list_evaluations_by_interview_id(interview["id"])
    report = generate_mock_report(interview, agents, messages, evaluations)
    existing = find_report_by_interview_id(interview["id"], user["id"])
    report_id = existing["id"] if existing else str(uuid4())

    if existing:
        response.status_code = 200
        db.execute(
            """
            UPDATE interview_reports
            SET total_score = ?, grade = ?, pass_recommendation = ?, ability_radar = ?,
                agent_feedback = ?, timeline_review = ?, summary = ?, suggestions = ?,
                updated_at = CURRENT_TIMESTAMP
            WHERE id = ? AND user_id = ?
            """,
            (
                report["total_score"],
                report["grade"],
                report["pass_recommendation"],
                json.dumps(report["ability_radar"], ensure_ascii=False),
                json.dumps(report["agent_feedback"], ensure_ascii=False),
                json.dumps(report["timeline_review"], ensure_ascii=False),
                report["summary"],
                report["suggestions"],
                report_id,
                user["id"],
            ),
        )
    else:
        db.execute(
            """
            INSERT INTO interview_reports (
              id, user_id, interview_id, total_score, grade, pass_recommendation,
              ability_radar, agent_feedback, timeline_review, summary, suggestions
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                report_id,
                user["id"],
                interview["id"],
                report["total_score"],
                report["grade"],
                report["pass_recommendation"],
                json.dumps(report["ability_radar"], ensure_ascii=False),
                json.dumps(report["agent_feedback"], ensure_ascii=False),
                json.dumps(report["timeline_review"], ensure_ascii=False),
                report["summary"],
                report["suggestions"],
            ),
        )
    db.commit()
    refresh_user_skill_stats(user["id"])
    return {"report": find_report_by_id(report_id, user["id"])}


@app.get("/api/interviews/{interview_id}/report")
def get_interview_report(interview_id: str, user: dict = Depends(require_auth)):
    interview = find_interview_by_user_id(interview_id, user["id"])
    if not interview:
        raise error(404, "面试不存在。")
    report = find_report_by_interview_id(interview["id"], user["id"])
    if not report:
        raise error(404, "报告不存在。")
    return {"report": report}


@app.get("/api/reports")
def list_reports(user: dict = Depends(require_auth)):
    return {"reports": list_reports_by_user_id(user["id"])}


@app.get("/api/reports/{report_id}")
def get_report(report_id: str, user: dict = Depends(require_auth)):
    report = find_report_by_id(report_id, user["id"])
    if not report:
        raise error(404, "报告不存在。")
    return {"report": report}


@app.get("/api/stats/me")
def get_stats(user: dict = Depends(require_auth)):
    return {"stats": refresh_user_skill_stats(user["id"])}


@app.get("/api/stats/me/dimensions")
def get_dimensions(user: dict = Depends(require_auth)):
    stats = refresh_user_skill_stats(user["id"])
    return {"dimensions": build_user_dimension_stats(stats)}
